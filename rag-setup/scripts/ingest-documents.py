#!/usr/bin/env python3
"""
Document ingestion script for Chimera RAG system.

Processes documents from the documents/ folder and prepares them for Khoj.
Supports: PDF, TXT, MD, DOCX, JSON

Usage:
    python ingest-documents.py
    python ingest-documents.py --folder /path/to/docs
    python ingest-documents.py --file document.pdf
"""

import os
import json
import glob
from pathlib import Path
from typing import List, Dict
import argparse

# Optional: pip install pypdf python-docx
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    print("Warning: pypdf not installed. PDF support disabled. Install with: pip install pypdf")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. DOCX support disabled. Install with: pip install python-docx")


class DocumentProcessor:
    """Process various document formats."""

    def __init__(self, documents_dir: str = "../documents"):
        self.documents_dir = Path(documents_dir)
        self.documents_dir.mkdir(exist_ok=True)

    def process_txt(self, filepath: Path) -> Dict:
        """Process plain text files."""
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            "filename": filepath.name,
            "content": content,
            "type": "text",
            "size": len(content)
        }

    def process_markdown(self, filepath: Path) -> Dict:
        """Process Markdown files."""
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            "filename": filepath.name,
            "content": content,
            "type": "markdown",
            "size": len(content)
        }

    def process_json(self, filepath: Path) -> Dict:
        """Process JSON files."""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        content = json.dumps(data, indent=2)
        return {
            "filename": filepath.name,
            "content": content,
            "type": "json",
            "size": len(content)
        }

    def process_pdf(self, filepath: Path) -> Dict:
        """Process PDF files."""
        if not HAS_PYPDF:
            return {
                "filename": filepath.name,
                "error": "pypdf not installed. Install with: pip install pypdf",
                "type": "pdf"
            }

        try:
            reader = PdfReader(filepath)
            content = "\n".join(page.extract_text() for page in reader.pages)
            return {
                "filename": filepath.name,
                "content": content,
                "type": "pdf",
                "pages": len(reader.pages),
                "size": len(content)
            }
        except Exception as e:
            return {
                "filename": filepath.name,
                "error": str(e),
                "type": "pdf"
            }

    def process_docx(self, filepath: Path) -> Dict:
        """Process DOCX files."""
        if not HAS_DOCX:
            return {
                "filename": filepath.name,
                "error": "python-docx not installed. Install with: pip install python-docx",
                "type": "docx"
            }

        try:
            doc = Document(filepath)
            content = "\n".join(para.text for para in doc.paragraphs)
            return {
                "filename": filepath.name,
                "content": content,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "size": len(content)
            }
        except Exception as e:
            return {
                "filename": filepath.name,
                "error": str(e),
                "type": "docx"
            }

    def process_file(self, filepath: Path) -> Dict:
        """Process a single file based on extension."""
        if not filepath.exists():
            return {"filename": filepath.name, "error": "File not found"}

        suffix = filepath.suffix.lower()

        if suffix == ".txt":
            return self.process_txt(filepath)
        elif suffix in [".md", ".markdown"]:
            return self.process_markdown(filepath)
        elif suffix == ".json":
            return self.process_json(filepath)
        elif suffix == ".pdf":
            return self.process_pdf(filepath)
        elif suffix == ".docx":
            return self.process_docx(filepath)
        else:
            return {
                "filename": filepath.name,
                "error": f"Unsupported file type: {suffix}",
                "supported": [".txt", ".md", ".json", ".pdf", ".docx"]
            }

    def process_directory(self, directory: Path = None) -> List[Dict]:
        """Process all documents in a directory."""
        if directory is None:
            directory = self.documents_dir

        directory = Path(directory)
        if not directory.exists():
            print(f"Directory not found: {directory}")
            return []

        results = []
        supported_extensions = ["*.txt", "*.md", "*.markdown", "*.json", "*.pdf", "*.docx"]

        for pattern in supported_extensions:
            for filepath in glob.glob(str(directory / f"**/{pattern}"), recursive=True):
                print(f"Processing: {filepath}")
                result = self.process_file(Path(filepath))
                results.append(result)

        return results

    def print_summary(self, results: List[Dict]):
        """Print processing summary."""
        print("\n" + "="*60)
        print("PROCESSING SUMMARY")
        print("="*60)

        success = [r for r in results if "content" in r]
        errors = [r for r in results if "error" in r]

        print(f"\nSuccessfully processed: {len(success)}")
        for result in success:
            print(f"  ✓ {result['filename']} ({result['type']}) - {result.get('size', 0)} chars")

        if errors:
            print(f"\nErrors: {len(errors)}")
            for error in errors:
                print(f"  ✗ {error['filename']} - {error['error']}")

        total_chars = sum(r.get('size', 0) for r in success)
        print(f"\nTotal characters: {total_chars:,}")
        print("="*60)


def main():
    parser = argparse.ArgumentParser(
        description="Ingest documents for Chimera RAG system"
    )
    parser.add_argument(
        "--folder",
        default="../documents",
        help="Documents folder path (default: ../documents)"
    )
    parser.add_argument(
        "--file",
        help="Process single file instead of folder"
    )

    args = parser.parse_args()

    processor = DocumentProcessor(args.folder)

    if args.file:
        # Process single file
        result = processor.process_file(Path(args.file))
        results = [result]
    else:
        # Process entire directory
        results = processor.process_directory()

    processor.print_summary(results)


if __name__ == "__main__":
    main()
